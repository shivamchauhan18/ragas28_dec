# =================================================
# 0️⃣ IMPORTS
# =================================================
import os, json, math, requests
import numpy as np
from dotenv import load_dotenv
from datasets import Dataset

# LangChain
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_ollama import OllamaEmbeddings, ChatOllama

# RAGAS
from ragas import evaluate
from ragas.metrics import (
    context_precision,
    context_recall,
    answer_relevancy,
    faithfulness
)

# NLP Metrics
import nltk
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from rouge_score import rouge_scorer
from nltk.translate.meteor_score import meteor_score
from bert_score import score as bert_score
from nltk.tokenize import word_tokenize, sent_tokenize
from textstat import flesch_reading_ease

# Reports
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm

nltk.download("punkt")
load_dotenv()


# =================================================
# 1️⃣ DOCUMENT
# =================================================
markdown = """
Deep learning is a subset of machine learning that uses neural networks with many layers.
It automatically learns hierarchical features like edges → shapes → objects.
It is used in NLP, computer vision, speech recognition, and generative AI.
"""


# =================================================
# 2️⃣ CHUNKING
# =================================================
splitter = RecursiveCharacterTextSplitter(chunk_size=80, chunk_overlap=10)
chunks = splitter.create_documents([markdown])

for i, c in enumerate(chunks):
    c.metadata["id"] = str(i)


# =================================================
# 3️⃣ VECTOR STORE
# =================================================
embeddings = OllamaEmbeddings(model="qwen3-embedding:0.6b")
vector_store = FAISS.from_documents(chunks, embeddings)
retriever = vector_store.as_retriever(search_kwargs={"k": 5})


# =================================================
# 4️⃣ LLM
# =================================================
llm = ChatOllama(
    model="mistral-small-3.1-24b:latest",
    temperature=0.2
)


# =================================================
# 5️⃣ PROMPT
# =================================================
prompt = PromptTemplate(
    template="""
Answer ONLY from the provided context.
If the answer is not found, say "I don't know."

Context:
{context}

Question:
{question}
""",
    input_variables=["context", "question"]
)


# =================================================
# 6️⃣ QUESTIONS & REFERENCES
# =================================================
questions = [
    "What is deep learning?",
    "How does deep learning learn hierarchical features?"
]

references = [
    "Deep learning uses neural networks with many layers.",
    "Deep learning learns hierarchical features through layered neural networks."
]


# =================================================
# 7️⃣ RAG PIPELINE
# =================================================
answers, contexts, retrieved_ids, retrieved_docs_all = [], [], [], []

for q in questions:
    docs = retriever.invoke(q)
    retrieved_ids.append([d.metadata["id"] for d in docs])

    doc_texts = [d.page_content for d in docs]
    retrieved_docs_all.append(doc_texts)

    ctx = "\n\n".join(doc_texts)
    final_prompt = prompt.invoke({"context": ctx, "question": q})

    ans = llm.invoke(final_prompt).content
    answers.append(ans)
    contexts.append(doc_texts)


# =================================================
# 8️⃣ IR METRICS
# =================================================
def hit_rate(y_true, y_pred):
    return sum(any(d in t for d in p) for t, p in zip(y_true, y_pred)) / len(y_true)

def mean_reciprocal_rank(y_true, y_pred):
    rr = []
    for t, p in zip(y_true, y_pred):
        score = 0
        for i, d in enumerate(p, 1):
            if d in t:
                score = 1 / i
                break
        rr.append(score)
    return float(np.mean(rr))

def ndcg(y_true, y_pred, k=5):
    scores = []
    for t, p in zip(y_true, y_pred):
        dcg = sum(1 / np.log2(i + 2) for i, d in enumerate(p[:k]) if d in t)
        idcg = sum(1 / np.log2(i + 2) for i in range(len(t[:k])))
        scores.append(dcg / idcg if idcg else 0)
    return float(np.mean(scores))

def recall_precision_at_k(y_true, y_pred, k=5):
    recall, precision = [], []
    for t, p in zip(y_true, y_pred):
        hits = sum(d in t for d in p[:k])
        recall.append(hits / len(t))
        precision.append(hits / k)
    return np.mean(recall), np.mean(precision)


relevant_ids = [c.metadata["id"] for c in chunks]
y_true = [relevant_ids for _ in questions]
y_pred = retrieved_ids


# =================================================
# 9️⃣ GENERATION METRICS
# =================================================
def compute_bleu(preds, refs):
    smooth = SmoothingFunction().method1
    return np.mean([
        sentence_bleu([word_tokenize(r)], word_tokenize(p), smoothing_function=smooth)
        for p, r in zip(preds, refs)
    ])

def compute_rouge(preds, refs):
    scorer = rouge_scorer.RougeScorer(["rouge1", "rougeL"], use_stemmer=True)
    r1, rL = [], []
    for p, r in zip(preds, refs):
        s = scorer.score(r, p)
        r1.append(s["rouge1"].fmeasure)
        rL.append(s["rougeL"].fmeasure)
    return np.mean(r1), np.mean(rL)

def compute_meteor(preds, refs):
    return np.mean([meteor_score([r], p) for p, r in zip(preds, refs)])

def compute_bertscore(preds, refs):
    _, _, f1 = bert_score(preds, refs, lang="en", rescale_with_baseline=True)
    return float(f1.mean())


# =================================================
# 🔟 CUSTOM RAG QUALITY METRICS (YOUR FUNCTIONS)
# =================================================
def answer_relevance_context_utilization(responses, references, retrieved_docs, top_k=5):
    rel, ctx = [], []
    for r, ref, docs in zip(responses, references, retrieved_docs):
        r_w, ref_w = set(word_tokenize(r.lower())), set(word_tokenize(ref.lower()))
        d_w = set(w for d in docs[:top_k] for w in word_tokenize(d.lower()))
        rel.append(len(r_w & ref_w) / len(ref_w))
        ctx.append(len(r_w & d_w) / len(r_w))
    return np.mean(rel), np.mean(ctx)

def groundedness(responses, retrieved_docs):
    return np.mean([
        len(set(word_tokenize(r.lower())) &
            set(w for d in docs for w in word_tokenize(d.lower()))) / len(word_tokenize(r))
        for r, docs in zip(responses, retrieved_docs)
    ])

def hallucination_rate(responses, retrieved_docs):
    return np.mean([
        len(set(word_tokenize(r.lower())) -
            set(w for d in docs for w in word_tokenize(d.lower()))) / len(word_tokenize(r))
        for r, docs in zip(responses, retrieved_docs)
    ])

def coherence_readability(responses):
    coh = [len(word_tokenize(r)) / len(sent_tokenize(r)) for r in responses]
    read = [flesch_reading_ease(r) for r in responses]
    return np.mean(coh), np.mean(read)

def query_relevancy(responses, queries):
    return np.mean([
        len(set(word_tokenize(r.lower())) &
            set(word_tokenize(q.lower()))) / len(word_tokenize(q))
        for r, q in zip(responses, queries)
    ])


# =================================================
# 1️⃣1️⃣ RAGAS
# =================================================
ragas_results = []

for i in range(len(questions)):
    ds = Dataset.from_dict({
        "question": [questions[i]],
        "answer": [answers[i]],
        "contexts": [contexts[i]],
        "reference": [references[i]]
    })

    res = evaluate(
        ds,
        metrics=[context_precision, context_recall, answer_relevancy, faithfulness],
        llm=llm,
        embeddings=embeddings
    )

    ragas_results.append(res.to_pandas().to_dict("records")[0])


# =================================================
# 1️⃣2️⃣ SAVE REPORTS
# =================================================
doc = Document()
doc.add_heading("RAG Evaluation Report", level=1)

doc.add_heading("Generation Metrics", level=2)
doc.add_paragraph(f"BLEU: {compute_bleu(answers, references):.4f}")
doc.add_paragraph(f"ROUGE-1 / ROUGE-L: {compute_rouge(answers, references)}")
doc.add_paragraph(f"METEOR: {compute_meteor(answers, references):.4f}")
doc.add_paragraph(f"BERTScore F1: {compute_bertscore(answers, references):.4f}")

doc.save("rag_evaluation_report.docx")

print("\n✅ RAG Evaluation Completed")
print("📄 DOCX: rag_evaluation_report.docx")
