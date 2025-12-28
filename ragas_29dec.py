# =================================================
# 0️⃣ IMPORTS
# =================================================
import numpy as np
import nltk
from typing import List
from dotenv import load_dotenv
from datasets import Dataset

# LangChain
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage
from langchain_ollama import OllamaEmbeddings, ChatOllama

# RAGAS (LEGACY – OLLAMA COMPATIBLE)
from ragas import evaluate
from ragas.metrics import (
    context_precision,
    context_recall,
    answer_relevancy,
    faithfulness,
)

# Generation / NLP Metrics
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from rouge_score import rouge_scorer
from bert_score import score as bert_score
from nltk.tokenize import word_tokenize, sent_tokenize
from textstat import flesch_reading_ease

# Report
from docx import Document

nltk.download("punkt")
load_dotenv()


# =================================================
# 1️⃣ DOCUMENT
# =================================================
markdown = """
Deep learning is a subset of machine learning that uses neural networks with many layers.
It automatically learns hierarchical features like edges → shapes → objects.
It is used in NLP, computer vision, speech recognition, and generative AI.
"""


# =================================================
# 2️⃣ CHUNKING
# =================================================
splitter = RecursiveCharacterTextSplitter(chunk_size=80, chunk_overlap=10)
chunks = splitter.create_documents([markdown])

for i, c in enumerate(chunks):
    c.metadata["id"] = str(i)


# =================================================
# 3️⃣ VECTOR STORE
# =================================================
embeddings = OllamaEmbeddings(model="nomic-embed-text:latest")
vector_store = FAISS.from_documents(chunks, embeddings)
retriever = vector_store.as_retriever(search_kwargs={"k": 20})


# =================================================
# 4️⃣ LLM
# =================================================
llm = ChatOllama(model="gemma3:1b", temperature=0.2)


# =================================================
# 5️⃣ PROMPT
# =================================================
prompt = PromptTemplate(
    template="""
Answer ONLY from the provided context.
If the answer is not found, say "I don't know."

Context:
{context}

Question:
{question}
""",
    input_variables=["context", "question"],
)


# =================================================
# 6️⃣ QUESTIONS & REFERENCES
# =================================================
questions = [
    "What is deep learning?",
    "How does deep learning learn hierarchical features?",
]

references = [
    "Deep learning uses neural networks with many layers.",
    "Deep learning learns hierarchical features through layered neural networks.",
]


# =================================================
# 7️⃣ RETRIEVAL + GENERATION
# =================================================
answers, contexts = [], []
retrieved_docs_per_q, retrieved_ids_per_q = [], []

for q in questions:
    docs = retriever.invoke(q)
    retrieved_docs_per_q.append(docs)
    retrieved_ids_per_q.append([d.metadata["id"] for d in docs])

    ctx = "\n\n".join(d.page_content for d in docs)
    final_prompt = prompt.format(context=ctx, question=q)

    ans = llm.invoke([HumanMessage(content=final_prompt)]).content
    answers.append(ans)
    contexts.append([d.page_content for d in docs])


# =================================================
# 8️⃣ STRICT GROUND TRUTH
# =================================================
def build_ground_truth_simple(references, retrieved_docs, default_id="0001"):
    gt = []
    for ref, docs in zip(references, retrieved_docs):
        ref_l = ref.lower()
        matches = [d.metadata["id"] for d in docs if ref_l in d.page_content.lower()]
        gt.append(matches if matches else [default_id])
    return gt


ground_truth_ids = build_ground_truth_simple(references, retrieved_docs_per_q)


# =================================================
# 9️⃣ IR METRICS
# =================================================
def hit_rate(y_true, y_pred):
    return sum(any(d in t for d in p) for t, p in zip(y_true, y_pred)) / len(y_true)

def mean_reciprocal_rank(y_true, y_pred):
    return np.mean([
        1 / (i + 1)
        for true, pred in zip(y_true, y_pred)
        for i, d in enumerate(pred)
        if d in true
    ] or [0])

def ndcg(y_true, y_pred, k=5):
    return np.mean([
        sum(1 / np.log2(i + 2) for i, d in enumerate(pred[:k]) if d in true)
        / max(1, sum(1 / np.log2(i + 2) for i in range(min(k, len(true)))))
        for true, pred in zip(y_true, y_pred)
    ])

def recall_precision_at_k(y_true, y_pred, k=5):
    rec, prec = [], []
    for t, p in zip(y_true, y_pred):
        hits = sum(d in t for d in p[:k])
        rec.append(hits / len(t))
        prec.append(hits / k)
    return np.mean(rec), np.mean(prec)


ir_metrics_per_q = []
k = 5
for i in range(len(questions)):
    y_true = [ground_truth_ids[i]]
    y_pred = [retrieved_ids_per_q[i]]
    r, p = recall_precision_at_k(y_true, y_pred, k)
    ir_metrics_per_q.append({
        f"Precision@{k}": p,
        f"Recall@{k}": r,
        f"HitRate@{k}": hit_rate(y_true, y_pred),
        f"MRR@{k}": mean_reciprocal_rank(y_true, y_pred),
        f"NDCG@{k}": ndcg(y_true, y_pred, k),
    })


# =================================================
# 🔟 CUSTOM RAG QUALITY METRICS
# =================================================
def answer_relevance_context_utilization(responses, references, docs):
    rel, ctx = [], []
    for r, ref, d in zip(responses, references, docs):
        rw, fw = set(word_tokenize(r.lower())), set(word_tokenize(ref.lower()))
        rel.append(len(rw & fw) / len(fw))
        dw = set(w for x in d for w in word_tokenize(x.page_content.lower()))
        ctx.append(len(rw & dw) / len(rw))
    return np.mean(rel), np.mean(ctx)

def groundedness(responses, docs):
    return np.mean([
        len(set(word_tokenize(r.lower())) &
            set(w for d in docs[i] for w in word_tokenize(d.page_content.lower())))
        / len(word_tokenize(r))
        for i, r in enumerate(responses)
    ])

def hallucination_rate(responses, docs):
    return np.mean([
        len(set(word_tokenize(r.lower())) -
            set(w for d in docs[i] for w in word_tokenize(d.page_content.lower())))
        / len(word_tokenize(r))
        for i, r in enumerate(responses)
    ])

def coherence_readability(responses):
    return (
        np.mean([len(word_tokenize(r)) / len(sent_tokenize(r)) for r in responses]),
        np.mean([flesch_reading_ease(r) for r in responses]),
    )

def query_relevancy(responses, queries):
    return np.mean([
        len(set(word_tokenize(r.lower())) & set(word_tokenize(q.lower())))
        / len(word_tokenize(q))
        for r, q in zip(responses, queries)
    ])


# =================================================
# 1️⃣1️⃣ RAGAS METRICS
# =================================================
ragas_results = []
for i in range(len(questions)):
    ds = Dataset.from_dict({
        "question": [questions[i]],
        "answer": [answers[i]],
        "contexts": [contexts[i]],
        "reference": [references[i]],
    })
    res = evaluate(
        ds,
        metrics=[context_precision, context_recall, answer_relevancy, faithfulness],
        llm=llm,
        embeddings=embeddings,
    )
    ragas_results.append(res.to_pandas().to_dict("records")[0])


# =================================================
# 1️⃣2️⃣ GENERATION METRICS (SAFE METEOR-LITE)
# =================================================
def compute_bleu(preds, refs):
    return np.mean([
        sentence_bleu([word_tokenize(r)], word_tokenize(p),
                      smoothing_function=SmoothingFunction().method1)
        for p, r in zip(preds, refs)
    ])

def compute_rouge(preds, refs):
    s = rouge_scorer.RougeScorer(["rouge1", "rougeL"], use_stemmer=True)
    r1, rL = [], []
    for p, r in zip(preds, refs):
        sc = s.score(r, p)
        r1.append(sc["rouge1"].fmeasure)
        rL.append(sc["rougeL"].fmeasure)
    return np.mean(r1), np.mean(rL)

def compute_meteor_lite(preds, refs):
    scores = []
    for p, r in zip(preds, refs):
        p_t, r_t = set(word_tokenize(p.lower())), set(word_tokenize(r.lower()))
        if not p_t or not r_t:
            scores.append(0)
            continue
        prec = len(p_t & r_t) / len(p_t)
        rec = len(p_t & r_t) / len(r_t)
        scores.append(2 * prec * rec / (prec + rec) if prec + rec else 0)
    return np.mean(scores)

def compute_bertscore(preds, refs):
    _, _, f1 = bert_score(preds, refs, lang="en", rescale_with_baseline=True)
    return float(f1.mean())


# =================================================
# 1️⃣3️⃣ SAVE DOCX REPORT
# =================================================
doc = Document()
doc.add_heading("RAG Evaluation Report", 1)

doc.add_heading("Generation Metrics", 2)
doc.add_paragraph(f"BLEU: {compute_bleu(answers, references):.4f}")
r1, rL = compute_rouge(answers, references)
doc.add_paragraph(f"ROUGE-1: {r1:.4f}")
doc.add_paragraph(f"ROUGE-L: {rL:.4f}")
doc.add_paragraph(f"METEOR (lite): {compute_meteor_lite(answers, references):.4f}")
doc.add_paragraph(f"BERTScore F1: {compute_bertscore(answers, references):.4f}")

doc.add_heading("Per-Question Metrics", 2)
for i, q in enumerate(questions):
    doc.add_heading(q, 3)
    doc.add_paragraph("IR Metrics:")
    for k, v in ir_metrics_per_q[i].items():
        doc.add_paragraph(f"{k}: {v:.4f}", style="List Continue")
    doc.add_paragraph("RAGAS Metrics:")
    for k, v in ragas_results[i].items():
        if isinstance(v, float):
            doc.add_paragraph(f"{k}: {v:.4f}", style="List Continue")

doc.add_heading("Aggregate RAG Quality Metrics", 2)
rel, ctx = answer_relevance_context_utilization(answers, references, retrieved_docs_per_q)
doc.add_paragraph(f"Answer Relevance: {rel:.4f}")
doc.add_paragraph(f"Context Utilization: {ctx:.4f}")
doc.add_paragraph(f"Groundedness: {groundedness(answers, retrieved_docs_per_q):.4f}")
doc.add_paragraph(f"Hallucination Rate: {hallucination_rate(answers, retrieved_docs_per_q):.4f}")
coh, read = coherence_readability(answers)
doc.add_paragraph(f"Coherence: {coh:.4f}")
doc.add_paragraph(f"Readability: {read:.4f}")
doc.add_paragraph(f"Query Relevancy: {query_relevancy(answers, questions):.4f}")

doc.save("rag_evaluation_report.docx")

print("\n✅ RAG Evaluation Completed Successfully")
